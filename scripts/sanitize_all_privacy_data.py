import os
import re
import json
import base64
import docx

DESKTOP_DIR = r'C:\Users\doctor\Desktop\Radiology_Templates_DOCX'
RADNITO_DIR = r'C:\Users\doctor\antigravity\adventurous-babbage\radnito'
RADIOLOGY_DICTATION_DIR = r'C:\Users\doctor\antigravity\adventurous-babbage\radiology-dictation'

# Privacy terms and names to scrub
SENSITIVE_PATTERNS = [
    (re.compile(r'\bKMCH\s+Stroke\s+Code\s+Proforma\b', re.IGNORECASE), 'Stroke Code Protocol'),
    (re.compile(r'\bKMCH\b', re.IGNORECASE), 'Hospital'),
    (re.compile(r'\bKovai\s+Medical\s+Center(\s+and\s+Hospital)?\b', re.IGNORECASE), 'Medical Center'),
    (re.compile(r'\b(Dr\.?|Doctor)\s+(Prakhyath|Gambira|Archive|Pragadesh|Mithun|Sathish|Suman|Suganya|Likhitha|Deepikha|Dharani|Abhirami|Bhuvanapriya|Gowtham|Santhosh|Sumathi|Arun|Sherene|Ibrahim)\b(\s+[A-Za-z\.]+)?', re.IGNORECASE), ''),
    (re.compile(r'\b(Co-read\s+by|Reported\s+by|Signed\s+by|Consultant\s+Radiologist|Radiologist:?)\s*:?.*', re.IGNORECASE), ''),
    (re.compile(r'\bDr\s+Prakhyath\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Mithun\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Suman\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Suganya\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Likhitha\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Deepikha\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Dharani\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Arun\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Pragadesh\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Gambira\b', re.IGNORECASE), ''),
    (re.compile(r'\bArchive\b', re.IGNORECASE), ''),
    (re.compile(r'\bPrakhyath\b', re.IGNORECASE), ''),
    (re.compile(r'\bGambira\b', re.IGNORECASE), ''),
    (re.compile(r'\bPragadesh\b', re.IGNORECASE), ''),
    (re.compile(r'\bMithun\b', re.IGNORECASE), ''),
    (re.compile(r'\bEMP\d+\b', re.IGNORECASE), ''),
    (re.compile(r'\\\\\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\\[^\s"\'\)]+', re.IGNORECASE), ''),
]

def sanitize_string(text):
    if not text:
        return text
    clean = text
    for pattern, replacement in SENSITIVE_PATTERNS:
        clean = pattern.sub(replacement, clean)
    # clean extra whitespace
    clean = re.sub(r'[ \t]+', ' ', clean).strip()
    return clean

def sanitize_docx_file(filepath):
    doc = docx.Document(filepath)
    changed = False
    
    # We want to remove paragraphs that become empty after sanitizing or only contain doctor names
    paragraphs_to_keep = []
    
    for p in doc.paragraphs:
        orig = p.text
        sanitized = sanitize_string(orig)
        if orig != sanitized:
            changed = True
            # clear runs and rewrite sanitized text
            for r in p.runs:
                r.text = ""
            if sanitized:
                p.add_run(sanitized)
        if p.text.strip():
            paragraphs_to_keep.append(p.text.strip())

    if changed:
        doc.save(filepath)

    # Return base64 and lines
    with open(filepath, 'rb') as f:
        b64 = base64.b64encode(f.read()).decode('utf-8')
    return b64, paragraphs_to_keep

def main():
    print("=================================================================")
    print("RUNNING THOROUGH PRIVACY AND SENSITIVE DATA SCRUB")
    print("=================================================================")

    # 1. Remove any unwanted multi-archive files from Desktop
    desktop_mri_a = os.path.join(DESKTOP_DIR, '02_MRI_A')
    unwanted_desktop_files = [
        'Archive formats.docx', 'MRI FORMAT.docx', 'MRI formats Sher.docx'
    ]
    for uf in unwanted_desktop_files:
        p = os.path.join(desktop_mri_a, uf)
        if os.path.exists(p):
            os.remove(p)
            print(f"[REMOVED UNWANTED FILE] {p}")

    # 2. Sanitize all remaining DOCX files on Desktop
    print("\n--- Sanitizing Desktop DOCX files ---")
    all_clean_templates = []
    global_id = 1

    for folder_name, category in [('01_CT', 'CT'), ('02_MRI_A', 'MRI A'), ('03_MRI_B', 'MRI B')]:
        folder_path = os.path.join(DESKTOP_DIR, folder_name)
        if not os.path.exists(folder_path):
            continue
        
        for f in sorted(os.listdir(folder_path)):
            if f.endswith('.docx') and not f.startswith('~$'):
                filepath = os.path.join(folder_path, f)
                b64, lines = sanitize_docx_file(filepath)
                if not lines:
                    continue
                
                title = sanitize_string(lines[0])
                if not title or len(title) > 90:
                    title = sanitize_string(os.path.splitext(f)[0])

                all_clean_templates.append({
                    'id': f"tpl_{global_id}",
                    'name': title,
                    'title': title,
                    'modality': category,
                    'category': category,
                    'lines': [sanitize_string(l) for l in lines if sanitize_string(l)],
                    'docxBase64': b64,
                    'content': '\n'.join([sanitize_string(l) for l in lines if sanitize_string(l)])
                })
                global_id += 1

    print(f"Total sanitized templates: {len(all_clean_templates)}")
    print(f"  CT: {len([t for t in all_clean_templates if t['category'] == 'CT'])}")
    print(f"  MRI A: {len([t for t in all_clean_templates if t['category'] == 'MRI A'])}")
    print(f"  MRI B: {len([t for t in all_clean_templates if t['category'] == 'MRI B'])}")

    # 3. Update templatesData.json in both projects
    for proj_dir in [RADNITO_DIR, RADIOLOGY_DICTATION_DIR]:
        for sub in ['services', 'public']:
            target = os.path.join(proj_dir, sub, 'templatesData.json')
            if os.path.exists(os.path.dirname(target)):
                with open(target, 'w', encoding='utf-8') as f:
                    json.dump(all_clean_templates, f, indent=2)
                print(f"[OK] Saved clean {target}")

    # 4. Sanitize report_knowledgebase.json in both projects
    for proj_dir in [RADNITO_DIR, RADIOLOGY_DICTATION_DIR]:
        for sub in ['services', 'public']:
            kb_target = os.path.join(proj_dir, sub, 'report_knowledgebase.json')
            if os.path.exists(kb_target):
                with open(kb_target, 'r', encoding='utf-8-sig') as f:
                    kb_data = json.load(f)
                
                # Sanitize content in all categories
                reports_by_cat = kb_data.get('reports_by_category', {})
                for cat, r_list in reports_by_cat.items():
                    for r in r_list:
                        if 'content' in r:
                            r['content'] = sanitize_string(r['content'])
                        if 'title' in r:
                            r['title'] = sanitize_string(r['title'])

                with open(kb_target, 'w', encoding='utf-8') as f:
                    json.dump(kb_data, f, indent=2)
                print(f"[OK] Sanitized {kb_target}")

    # 5. Clean scripts folder from any raw network paths or sensitive names
    scripts_dir = os.path.join(RADNITO_DIR, 'scripts')
    for sf in os.listdir(scripts_dir):
        if sf.endswith('.py'):
            sp = os.path.join(scripts_dir, sf)
            with open(sp, 'r', encoding='utf-8') as f:
                content = f.read()
            # Replace internal network UNC path with placeholder
            clean_code = re.sub(r'\\\\172\.16\.1\.85\\[^\s"\'\)]+', r'[NETWORK_RADIOLOGY_SHARE]', content)
            clean_code = re.sub(r'MRI_Source', 'MRI_Source', clean_code)
            clean_code = re.sub(r'CT_Source', 'CT_Source', clean_code)
            clean_code = re.sub(r'CT_Specialized', 'CT_Specialized', clean_code)
            clean_code = re.sub(r'Archive', 'Archive', clean_code)
            if clean_code != content:
                with open(sp, 'w', encoding='utf-8') as f:
                    f.write(clean_code)
                print(f"[CLEANED CODE] {sp}")

    print("\nSanitation completed successfully!")

if __name__ == '__main__':
    main()
