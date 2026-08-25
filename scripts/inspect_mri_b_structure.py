import docx
import zipfile
import xml.etree.ElementTree as ET

doc_path = r'D:\Desktop\MRI B.docx'

with zipfile.ZipFile(doc_path) as z:
    xml_content = z.read('word/document.xml')

root = ET.fromstring(xml_content)
namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# Find all paragraphs and analyze how titles / breaks appear
doc = docx.Document(doc_path)
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total sections: {len(doc.sections)}")
print(f"Total tables: {len(doc.tables)}")

print("\n--- ALL PARAGRAPHS WITH POTENTIAL TITLES (ALL CAPS OR HEADING STYLES) ---")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    # Check if text is all caps or starts with MRI or SCAN or BRAIN or SPINE or looks like title
    is_caps = text.isupper() and len(text) > 4
    has_mri = 'MRI' in text.upper() or 'SCAN' in text.upper()
    runs_xml = "".join(r._r.xml for r in p.runs)
    has_br = 'w:br' in runs_xml or 'w:pageBreakBefore' in p._p.xml
    
    if is_caps or has_mri or has_br or i == 0 or 'TECHNIQUE' in text.upper() or 'IMPRESSION' in text.upper():
        print(f"P#{i:03d} [Break:{has_br}] [Style:{p.style.name}]: {text[:80]}")
