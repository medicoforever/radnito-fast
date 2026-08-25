import docx
import os

ct_prag_path = r'[NETWORK_RADIOLOGY_SHARE] Pragadesh\FORMATS\Radiology report formats\CT reporting formats.docx'
doc = docx.Document(ct_prag_path)
print(f"Paragraphs in CT reporting formats.docx: {len(doc.paragraphs)}")

for i in range(min(30, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    if p.text.strip():
        print(f"P{i:02d}: {p.text[:80]}")
