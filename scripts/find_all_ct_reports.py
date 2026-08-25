import docx
import re

ct_path = r'\\172.16.1.85\Radiology\Dr Pragadesh\FORMATS\Radiology report formats\CT reporting formats.docx'
doc = docx.Document(ct_path)

print(f"Total paragraphs in CT reporting formats.docx: {len(doc.paragraphs)}")

report_candidates = []

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if not txt:
        continue
    is_bold = any(r.bold for r in p.runs)
    is_under = any(r.underline for r in p.runs)
    
    # Check if this paragraph is a title of a new CT report
    # Typically bold or underlined, and contains scan name / anatomy / protocol
    # Not 'IMPRESSION:', not 'CLINICAL PROFILE:', not 'FINDINGS:', not 'TECHNIQUE:'
    is_section_header = any(txt.upper().startswith(h) for h in [
        'IMPRESSION', 'CLINICAL', 'PAST HISTORY', 'TECHNIQUE', 'FINDINGS', 
        'OBSERVATIONS', 'PROTOCOL', 'ADVICE', 'RECOMMENDATION', 'REMARKS',
        'CORADS', 'CALCIUM SCORE', 'DOMINANCE'
    ])
    
    if not is_section_header and (is_bold or is_under or txt.isupper()):
        if any(k in txt.upper() for k in ['C.T.', 'CT ', 'HRCT', 'SCAN', 'ANGIO', 'STROKE CODE', 'MYELOGRAM', 'ORBIT', 'TEMPORAL', 'PNS', 'LUNGS', 'CARDIAC', 'TRAUMA', 'KUB', 'AORTA']):
            if len(txt) < 85:
                report_candidates.append((i, txt, is_bold, is_under))

print(f"Found {len(report_candidates)} potential report starts:")
for idx, (p_idx, title, b, u) in enumerate(report_candidates):
    print(f"#{idx+1:02d} P#{p_idx:04d} [B={b}, U={u}]: {title}")
