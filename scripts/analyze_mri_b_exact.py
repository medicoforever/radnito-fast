import docx

doc = docx.Document(r'D:\Desktop\MRI B.docx')

print(f"Total paragraphs in MRI B.docx: {len(doc.paragraphs)}")

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    is_bold = any(r.bold for r in p.runs)
    is_under = any(r.underline for r in p.runs)
    is_italic = any(r.italic for r in p.runs)
    
    # Print potential headings or report titles
    if txt and (is_bold or is_under or txt.upper().startswith('MRI') or txt.upper().startswith('BRACHIAL') or 'SCAN OF' in txt.upper() or 'IMPRESSION' in txt.upper()):
        print(f"P#{i:03d} [B={is_bold}, U={is_under}, I={is_italic}]: {txt[:75]}")
