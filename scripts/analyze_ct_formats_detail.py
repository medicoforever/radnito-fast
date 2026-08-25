import docx

ct_path = r'\\172.16.1.85\Radiology\Dr Pragadesh\FORMATS\Radiology report formats\CT reporting formats.docx'
doc = docx.Document(ct_path)

print(f"Total paragraphs in CT reporting formats.docx: {len(doc.paragraphs)}")

for i in range(1090, 1280):
    p = doc.paragraphs[i]
    txt = p.text.strip()
    is_bold = any(r.bold for r in p.runs)
    is_under = any(r.underline for r in p.runs)
    if txt and (is_bold or is_under or txt.isupper() or 'SCAN' in txt.upper() or 'STROKE' in txt.upper() or 'CARDIAC' in txt.upper() or 'IMPRESSION' in txt.upper() or 'TECHNIQUE' in txt.upper()):
        print(f"P#{i:04d} [B={is_bold}, U={is_under}]: {txt[:75]}")
