import os
import re
import json
import base64
import copy
import shutil
import docx
from docx.shared import Pt, Inches

DESKTOP_DIR = r'C:\Users\doctor\Desktop\Radiology_Templates_DOCX'
RADNITO_DIR = r'C:\Users\doctor\antigravity\adventurous-babbage\radnito'
RADIOLOGY_DICTATION_DIR = r'C:\Users\doctor\antigravity\adventurous-babbage\radiology-dictation'

CT_MASTER_PATH = r'\\172.16.1.85\Radiology\Dr Pragadesh\FORMATS\Radiology report formats\CT reporting formats.docx'
CT_IHSR_PATH = r'\\172.16.1.85\Radiology\IHSR_New imaging Format templates\IHSR Format\CT IHSR'
MRI_A_PATH = r'\\172.16.1.85\Radiology\Dr Prakhyath Gambira\mri formats'
MRI_B_PATH = r'D:\Desktop\MRI B.docx'

SENSITIVE_PATTERNS = [
    (re.compile(r'\bKMCH\s+Stroke\s+Code\s+Proforma\b', re.IGNORECASE), 'Stroke Code Protocol'),
    (re.compile(r'\bKMCH\b', re.IGNORECASE), ''),
    (re.compile(r'\bKovai\s+Medical\s+Center(\s+and\s+Hospital)?\b', re.IGNORECASE), ''),
    (re.compile(r'\b(Dr\.?|Doctor)\s+(Prakhyath|Gambira|Gopinath|Pragadesh|Mithun|Sathish|Suman|Suganya|Likhitha|Deepikha|Dharani|Abhirami|Bhuvanapriya|Gowtham|Santhosh|Sumathi|Arun|Sherene|Ibrahim)\b(\s+[A-Za-z\.]+)?', re.IGNORECASE), ''),
    (re.compile(r'\b(Co-read\s+by|Reported\s+by|Signed\s+by|Consultant\s+Radiologist|Radiologist:?)\s*:?.*', re.IGNORECASE), ''),
    (re.compile(r'\bDr\s+Prakhyath\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Mithun\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Suman\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Suganya\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Likhitha\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Deepikha\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Dharani\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Arun\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Pragadesh\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Gambira\b', re.IGNORECASE), ''),
    (re.compile(r'\bGopinath\b', re.IGNORECASE), ''),
    (re.compile(r'\bPrakhyath\b', re.IGNORECASE), ''),
    (re.compile(r'\bGambira\b', re.IGNORECASE), ''),
    (re.compile(r'\bPragadesh\b', re.IGNORECASE), ''),
    (re.compile(r'\bMithun\b', re.IGNORECASE), ''),
    (re.compile(r'\bEMP\d+\b', re.IGNORECASE), ''),
]

def sanitize_text(text):
    if not text:
        return text
    clean = text
    for pat, repl in SENSITIVE_PATTERNS:
        clean = pat.sub(repl, clean)
    clean = re.sub(r'[ \t]+', ' ', clean).strip()
    return clean

def sanitize_filename(name):
    clean = sanitize_text(name)
    clean = re.sub(r'[\\/*?:"<>|]', '_', clean)
    clean = re.sub(r'\s+', ' ', clean).strip()
    return clean[:90]

def create_cloned_doc_from_paragraphs(master_doc, start_idx, end_idx, target_path):
    new_doc = docx.Document()
    for s_master, s_new in zip(master_doc.sections, new_doc.sections):
        s_new.top_margin = s_master.top_margin
        s_new.bottom_margin = s_master.bottom_margin
        s_new.left_margin = s_master.left_margin
        s_new.right_margin = s_master.right_margin
        s_new.page_width = s_master.page_width
        s_new.page_height = s_master.page_height

    if len(new_doc.paragraphs) > 0:
        p_elem = new_doc.paragraphs[0]._p
        p_elem.getparent().remove(p_elem)

    body = new_doc._body._body
    for i in range(start_idx, end_idx):
        orig_p = master_doc.paragraphs[i]
        p_copy = copy.deepcopy(orig_p._p)
        
        # Sanitize text inside <w:t> tags while preserving all formatting (<w:b>, <w:i>, <w:u>)
        for t_elem in p_copy.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t_elem.text:
                for pat, repl in SENSITIVE_PATTERNS:
                    t_elem.text = pat.sub(repl, t_elem.text)
                    
        body.append(p_copy)

    # Ensure Times New Roman font & 12pt size across all runs
    for p in new_doc.paragraphs:
        for r in p.runs:
            r.font.name = 'Times New Roman'
            if r.font.size is None:
                r.font.size = Pt(12)

    new_doc.save(target_path)

def sanitize_single_docx_file(filepath):
    doc = docx.Document(filepath)
    changed = False
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = 'Times New Roman'
            if r.font.size is None:
                r.font.size = Pt(12)
            orig = r.text
            sanitized = orig
            for pat, repl in SENSITIVE_PATTERNS:
                sanitized = pat.sub(repl, sanitized)
            if orig != sanitized:
                r.text = sanitized
                changed = True
    doc.save(filepath)

def get_docx_base64_and_lines(filepath):
    doc = docx.Document(filepath)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    with open(filepath, 'rb') as f:
        b64 = base64.b64encode(f.read()).decode('utf-8')
    return b64, lines

def main():
    print("=================================================================")
    print("BUILDING 100% AUTHENTIC FORMATTED RADIOLOGY TEMPLATES")
    print("  (Times New Roman 12, Exact Bold, Underlines, Italics & Gaps)")
    print("=================================================================")

    # Clear desktop folder
    if os.path.exists(DESKTOP_DIR):
        shutil.rmtree(DESKTOP_DIR)
    os.makedirs(DESKTOP_DIR, exist_ok=True)

    ct_dir = os.path.join(DESKTOP_DIR, '01_CT')
    mri_a_dir = os.path.join(DESKTOP_DIR, '02_MRI_A')
    mri_b_dir = os.path.join(DESKTOP_DIR, '03_MRI_B')

    os.makedirs(ct_dir, exist_ok=True)
    os.makedirs(mri_a_dir, exist_ok=True)
    os.makedirs(mri_b_dir, exist_ok=True)

    all_catalog_templates = []
    global_id = 1

    # -----------------------------------------------------------------
    # 1. PROCESS CT (Centricity RIS-i Exact Formats)
    # -----------------------------------------------------------------
    print("\n--- 1. Processing CT Formats (Centricity RIS-i) ---")
    ct_seen = set()
    ct_doc = docx.Document(CT_MASTER_PATH)

    ct_boundaries = [
        (1, 26, "C.T.SCAN OF BRAIN (PLAIN)"),
        (26, 57, "C.T.SCAN OF BOTH TEMPORAL BONES"),
        (57, 77, "CT SCAN LUMBAR MYELOGRAM"),
        (77, 118, "C.T. SCAN OF TRAUMA BRAIN AND CERVICAL SPINE"),
        (118, 149, "CT SCAN OF BRAIN AND PNS"),
        (149, 178, "C.T.SCAN OF BRAIN - PLAIN AND CONTRAST"),
        (178, 205, "CT SCAN OF NECK"),
        (205, 245, "C.T. SCAN OF ABDOMEN (CONTRAST)"),
        (245, 283, "C.T. SCAN OF ABDOMEN (CONTRAST - DETAILED)"),
        (283, 322, "C.T. SCAN OF ABDOMEN (PLAIN - MALE)"),
        (322, 363, "C.T. SCAN OF ABDOMEN (PLAIN - FEMALE)"),
        (363, 391, "CT SCAN OF KIDNEYS (PLAIN)"),
        (391, 421, "C.T.RENAL ANGIO FOLLOWED BY IVP"),
        (421, 432, "CT WHOLE BODY TRAUMA"),
        (432, 451, "C.T. SCAN OF BRAIN"),
        (451, 470, "C.T.SCAN OF NECK"),
        (470, 488, "C.T.SCAN OF THORAX"),
        (488, 519, "C.T. SCAN OF ABDOMEN"),
        (519, 537, "CT SCAN OF WHOLE SPINE"),
        (537, 567, "CT ANGIOGRAPHY OF AORTA"),
        (567, 599, "CT ANGIOGRAPHY OF ARCH OF AORTA & THORACIC AORTA"),
        (599, 629, "BIFURCATION AORTAGRAPHY WITH RUN OFF"),
        (629, 663, "CT ANGIOGRAPHY OF BOTH LOWER LIMBS"),
        (663, 696, "CT ANGIOGRAPHY OF BRAIN"),
        (696, 719, "C.T.SCAN OF BRAIN AND ORBIT"),
        (719, 754, "CT ANGIO OF NECK"),
        (754, 782, "MULTISLICE CT PULMONARY ANGIOGRAPHY"),
        (782, 817, "CT ANGIO OF ABDOMEN"),
        (817, 858, "CT ANGIOGRAPHY OF BOTH LOWER LIMBS (RUN OFF)"),
        (858, 929, "CT CORONARY ANGIOGRAM"),
        (929, 971, "CT ANGIOGRAPHY OF THE PULMONARY ARTERIES"),
        (971, 1068, "CT ANGIOGRAM STUDY OF AORTA AND PULMONARY ARTERIES"),
        (1068, 1099, "HRCT SCAN OF LUNGS"),
        (1099, 1126, "C.T. SCAN OF BRAIN AND PITUITARY GLAND"),
        (1126, 1194, "CT Brain (Plain) - Stroke Code Protocol"),
        (1194, 1216, "CT CARDIAC SCREENING"),
        (1216, 1244, "C.T. SCAN OF FACIAL BONES INCLUDING MANDIBLE"),
        (1244, 1273, "HRCT SCAN OF LUNGS (COVID SCREENING)"),
        (1273, 1373, "CT CORONARY ANGIOGRAM AND AORTOGRAM (TAVI PROTOCOL)"),
        (1373, 1435, "POST CABG CT ANGIOGRAM"),
        (1435, 1449, "C.T.SCAN OF BRAIN (FOLLOW UP)"),
        (1449, 1476, "HRCT SCAN OF LUNGS (PRONE / SUPINE)"),
        (1476, 1493, "C.T.SCAN OF NECK (CONTRAST)"),
        (1493, 1511, "C.T.SCAN OF ORBITS"),
        (1511, 1548, "CT PNS"),
        (1548, 1571, "C.T.SCAN OF THORAX"),
        (1571, 1588, "C.T.SCAN OF UPPER ABDOMEN"),
        (1588, 1614, "HIGH RESOLUTION C.T. SCAN OF TEMPORAL BONES"),
        (1614, 1648, "CT SCAN OF KUB (PLAIN)"),
    ]

    for idx, (s_idx, e_idx, report_title) in enumerate(ct_boundaries):
        norm = re.sub(r'[^a-zA-Z0-9]', '', report_title.lower())
        ct_seen.add(norm)
        fname = f"{idx+1:02d}_{sanitize_filename(report_title)}.docx"
        fpath = os.path.join(ct_dir, fname)
        create_cloned_doc_from_paragraphs(ct_doc, s_idx, e_idx, fpath)
        b64, lines = get_docx_base64_and_lines(fpath)
        title = lines[0] if lines else report_title
        all_catalog_templates.append({
            'id': f"ct_{global_id}",
            'name': title,
            'title': title,
            'modality': 'CT',
            'category': 'CT',
            'lines': lines,
            'docxBase64': b64,
            'content': '\n'.join(lines)
        })
        global_id += 1
        print(f"  [CT #{idx+1:02d}] {report_title} ({e_idx - s_idx} paragraphs)")

    # -----------------------------------------------------------------
    # 2. PROCESS MRI A (\\172.16.1.85\Radiology\Dr Prakhyath Gambira\mri formats)
    # -----------------------------------------------------------------
    print("\n--- 2. Processing MRI A Formats (Root Individual Files) ---")
    ignored_mri_a = {
        'gopinath formats', 'dr prakhyath gambira - shortcut.lnk', 'sathish-march ppt.pptx',
        'gopinath formats - shortcut.lnk', 'fistula.png', 'gopinath formats.docx',
        'mri format.docx', 'mri formats sher.docx'
    }

    for item in sorted(os.listdir(MRI_A_PATH)):
        if item.lower() in ignored_mri_a or item.startswith('~$'):
            continue
        full_path = os.path.join(MRI_A_PATH, item)
        if os.path.isdir(full_path) or not item.endswith('.docx'):
            continue
        
        try:
            dst_path = os.path.join(mri_a_dir, item)
            shutil.copy2(full_path, dst_path)
            sanitize_single_docx_file(dst_path)
            b64, lines = get_docx_base64_and_lines(dst_path)
            if not lines:
                continue
            title = lines[0] if len(lines[0]) < 80 and (lines[0].isupper() or lines[0].startswith('MRI') or lines[0].startswith('BRACHIAL') or lines[0].startswith('CARDIAC')) else os.path.splitext(item)[0]
            all_catalog_templates.append({
                'id': f"mri_a_{global_id}",
                'name': title,
                'title': title,
                'modality': 'MRI A',
                'category': 'MRI A',
                'lines': lines,
                'docxBase64': b64,
                'content': '\n'.join(lines)
            })
            global_id += 1
            print(f"  [MRI A] {item} -> {title[:60]}")
        except Exception as e:
            print(f"  [MRI A Warning] {item}: {e}")

    # -----------------------------------------------------------------
    # 3. PROCESS MRI B (D:\Desktop\MRI B.docx)
    # -----------------------------------------------------------------
    print("\n--- 3. Processing MRI B Formats (D:\\Desktop\\MRI B.docx) ---")
    mri_b_doc = docx.Document(MRI_B_PATH)

    mri_b_boundaries = [
        (0, 44, "MRI SCAN OF CERVICAL SPINE AND SCREENING OF REST OF THE SPINE"),
        (44, 89, "MRI SCAN OF BRAIN"),
        (89, 136, "MRI SCAN OF LUMBOSACRAL SPINE AND SCREENING OF REST OF THE SPINE"),
        (136, 184, "MRI SCAN OF DORSAL SPINE AND SCREENING OF REST OF THE SPINE"),
        (184, 232, "MRI SCAN OF KNEE JOINT"),
        (232, 275, "MRI SCAN OF HIP JOINTS"),
        (275, 283, "MRI SCAN OF WRIST JOINT"),
        (283, 305, "MRI SCAN OF LEFT FOREARM"),
        (305, 348, "MRI BONY PELVIS WITH LUMBAR SPINE SCREENING"),
        (348, 384, "MRI SCAN OF BRAIN & ORBITS"),
        (384, 417, "MRI SCAN OF PELVIS"),
        (417, 451, "MRI SCAN OF NECK"),
        (451, 476, "BRACHIAL PLEXUS SCREENING"),
        (476, 512, "MRI SCAN OF LEG"),
        (512, 523, "MRI SCAN OF LEFT SHOULDER JOINT"),
        (523, 533, "MRI - MRCP"),
        (533, 541, "MRI PITUITARY GLAND"),
        (541, 564, "MRI SCAN OF RIGHT FOOT"),
    ]

    for idx, (s_idx, e_idx, report_title) in enumerate(mri_b_boundaries):
        fname = f"{idx+1:02d}_{sanitize_filename(report_title)}.docx"
        fpath = os.path.join(mri_b_dir, fname)
        create_cloned_doc_from_paragraphs(mri_b_doc, s_idx, e_idx, fpath)
        b64, lines = get_docx_base64_and_lines(fpath)
        title = lines[0] if lines else report_title
        all_catalog_templates.append({
            'id': f"mri_b_{global_id}",
            'name': title,
            'title': title,
            'modality': 'MRI B',
            'category': 'MRI B',
            'lines': lines,
            'docxBase64': b64,
            'content': '\n'.join(lines)
        })
        global_id += 1
        print(f"  [MRI B #{idx+1:02d}] {report_title} ({e_idx - s_idx} paragraphs)")

    # -----------------------------------------------------------------
    # 4. SUMMARY & SAVE TO PROJECTS
    # -----------------------------------------------------------------
    print("\n" + "=" * 65)
    print(f"TOTAL AUTHENTIC TEMPLATES ASSEMBLED: {len(all_catalog_templates)}")
    ct_count = len([t for t in all_catalog_templates if t['category'] == 'CT'])
    mri_a_count = len([t for t in all_catalog_templates if t['category'] == 'MRI A'])
    mri_b_count = len([t for t in all_catalog_templates if t['category'] == 'MRI B'])

    print(f"  [CT]:    {ct_count} templates")
    print(f"  [MRI A]: {mri_a_count} templates")
    print(f"  [MRI B]: {mri_b_count} templates")
    print("=" * 65)

    # Write README
    readme_path = os.path.join(DESKTOP_DIR, 'README_AND_INSTRUCTIONS.txt')
    with open(readme_path, 'w', encoding='utf-8') as f:
        f.write(f"""================================================================================
AUTHENTIC RADIOLOGY DOCX TEMPLATES (TIMES NEW ROMAN 12, EXACT STYLES & SPACING)
================================================================================
Workspace Location: C:\\Users\\doctor\\Desktop\\Radiology_Templates_DOCX

Total Active Templates: {len(all_catalog_templates)}
  📁 01_CT     ({ct_count} templates: GE Centricity RIS-i isolated CT reports)
  📁 02_MRI_A  ({mri_a_count} templates: Dr Prakhyath Gambira individual MRI formats)
  📁 03_MRI_B  ({mri_b_count} templates: MRI B master formats)

HOW TO USE:
--------------------------------------------------------------------------------
1. EDIT ANY TEMPLATE:
   - Double-click any .docx file to open and edit in Microsoft Word.
   - All underlines, italics, bolds, blank lines, and Times New Roman 12 are 100% preserved.
   - Save the file (Ctrl+S).

2. ADD NEW TEMPLATES:
   - Paste any new .docx into 01_CT, 02_MRI_A, or 03_MRI_B.

3. SYNC TO RADNITO & RADIOLOGY DICTATION APPS (1-CLICK):
   - Double-click "Sync_Templates_To_App.bat" in this folder.
   - It will automatically package and push everything to GitHub and GitHub Pages!
================================================================================
""")

    # Write sync script and batch file to Desktop
    shutil.copy2(os.path.join(RADNITO_DIR, 'scripts', 'sync_templates_from_desktop.py'), os.path.join(DESKTOP_DIR, 'sync_templates.py'))
    with open(os.path.join(DESKTOP_DIR, 'Sync_Templates_To_App.bat'), 'w') as f:
        f.write('@echo off\ntitle Sync Templates to Apps\npython "%~dp0sync_templates.py"\npause\n')

    # Update radnito & radiology-dictation templatesData.json
    for project_dir in [RADNITO_DIR, RADIOLOGY_DICTATION_DIR]:
        services_dir = os.path.join(project_dir, 'services')
        public_dir = os.path.join(project_dir, 'public')
        os.makedirs(services_dir, exist_ok=True)
        os.makedirs(public_dir, exist_ok=True)

        for p_target in [os.path.join(services_dir, 'templatesData.json'), os.path.join(public_dir, 'templatesData.json')]:
            with open(p_target, 'w', encoding='utf-8') as f:
                json.dump(all_catalog_templates, f, indent=2)
            print(f"[OK] Saved {p_target}")

if __name__ == '__main__':
    main()
