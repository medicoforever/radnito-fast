import docx
import re

ct_path = r'\\172.16.1.85\Radiology\Dr Pragadesh\FORMATS\Radiology report formats\CT reporting formats.docx'
doc = docx.Document(ct_path)

# Let's define the exact start indices of all reports in CT reporting formats.docx
# We know paragraph 0 to 1647
raw_reports = []

# List of known exact start headers / patterns
is_start_indices = [
    1,    # C.T.SCAN OF BRAIN (PLAIN)
    26,   # C.T.SCAN OF BOTH TEMPORAL BONES
    57,   # CT SCAN LUMBAR MYELOGRAM
    77,   # C.T. SCAN OF TRAUMA BRAIN AND CERVICAL SPINE
    118,  # CT SCAN OF BRAIN AND PNS
    149,  # C.T.SCAN OF BRAIN – PLAIN AND CONTRAST
    178,  # CT SCAN OF NECK
    205,  # C.T. SCAN OF ABDOMEN (CONTRAST)
    245,  # C.T. SCAN OF ABDOMEN (CONTRAST 2)
    283,  # C.T. SCAN OF ABDOMEN (PLAIN)
    322,  # C.T. SCAN OF ABDOMEN (PLAIN 2)
    363,  # CT SCAN OF KIDNEYS (PLAIN)
    391,  # C.T.RENAL ANGIO FOLLOWED BY IVP
    421,  # CT WHOLE BODY TRAUMA
    432,  # C.T. SCAN OF BRAIN
    451,  # C.T.SCAN OF NECK
    470,  # C.T.SCAN OF THORAX
    488,  # C.T. SCAN OF ABDOMEN
    519,  # CT SCAN OF WHOLE SPINE
    537,  # CT ANGIOGRAPHY OF AORTA
    567,  # CT ANGIOGRAPHY OF ARCH OF AORTA & THORACIC AORTA
    599,  # BIFURCATION AORTAGRAPHY WITH RUN OFF
    629,  # CT ANGIOGRAPHY OF BOTH LOWER LIMBS
    663,  # CT ANGIOGRAPHY OF BRAIN
    696,  # C.T.SCAN OF BRAIN AND ORBIT
    719,  # CT ANGIO OF NECK
    754,  # MULTISLICE CT PULMONARY ANGIOGRAPHY
    782,  # CT ANGIO OF ABDOMEN
    817,  # CT ANGIOGRAPHY OF BOTH LOWER LIMBS
    858,  # CT CORONARY ANGIOGRAM
    929,  # CT ANGIOGRAPHY OF THE PULMONARY ARTERIES
    971,  # CT ANGIOGRAM STUDY OF AORTA AND PULMONARY ARTERIES
    1068, # HRCT SCAN OF LUNGS
    1099, # C.T. SCAN OF BRAIN AND PITUITARY GLAND
    1126, # CT Brain (Plain) - Stroke Code Protocol
    1194, # CT CARDIAC SCREENING
    1216, # C.T. SCAN OF FACIAL BONES INCLUDING MANDIBLE
    1244, # HRCT SCAN OF LUNGS (COVID SCREENING)
    1273, # CT CORONARY ANGIOGRAM AND AORTOGRAM (TAVI PROTOCOL)
    1373, # POST CABG CT ANGIOGRAM
    1435, # C.T.SCAN OF BRAIN (FOLLOW UP)
    1449, # HRCT SCAN OF LUNGS
    1476, # C.T.SCAN OF NECK
    1493, # C.T.SCAN OF ORBITS
    1511, # CT PNS
    1548, # C.T.SCAN OF THORAX
    1571, # C.T.SCAN OF UPPER ABDOMEN
    1588, # HIGH RESOLUTION C.T. SCAN OF TEMPORAL BONES
    1614, # CT SCAN OF KUB (PLAIN)
]

for idx in range(len(is_start_indices)):
    start = is_start_indices[idx]
    end = is_start_indices[idx+1] if idx+1 < len(is_start_indices) else len(doc.paragraphs)
    title = doc.paragraphs[start].text.strip()
    if not title:
        title = doc.paragraphs[start+1].text.strip()
    print(f"CT Report #{idx+1:02d} [P#{start:04d}..P#{end:04d}, {end-start:2d} paras]: {title[:60]}")
