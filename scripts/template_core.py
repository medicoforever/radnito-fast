import copy
import os
import re
import json
import base64
import docx
from docx.shared import Pt, Inches

DESKTOP_DIR = r'C:\Users\doctor\Desktop\Radiology_Templates_DOCX'
RADNITO_DIR = r'C:\Users\doctor\antigravity\adventurous-babbage\radnito'
RADIOLOGY_DICTATION_DIR = r'C:\Users\doctor\antigravity\adventurous-babbage\radiology-dictation'

SENSITIVE_PATTERNS = [
    (re.compile(r'\bKMCH\s+Stroke\s+Code\s+Proforma\b', re.IGNORECASE), 'Stroke Code Protocol'),
    (re.compile(r'\bKMCH\b', re.IGNORECASE), ''),
    (re.compile(r'\bKovai\s+Medical\s+Center(\s+and\s+Hospital)?\b', re.IGNORECASE), ''),
    (re.compile(r'\b(Dr\.?|Doctor)\s+(Prakhyath|Gambira|Gopinath|Pragadesh|Mithun|Sathish|Suman|Suganya|Likhitha|Deepikha|Dharani|Abhirami|Bhuvanapriya|Gowtham|Santhosh|Sumathi|Arun|Sherene|Ibrahim)\b(\s+[A-Za-z\.]+)?', re.IGNORECASE), ''),
    (re.compile(r'\b(Co-read\s+by|Reported\s+by|Signed\s+by|Consultant\s+Radiologist|Radiologist:?)\s*:?.*', re.IGNORECASE), ''),
    (re.compile(r'\bDr\s+Prakhyath\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Mithun\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Suman\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Suganya\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Likhitha\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Deepikha\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Dharani\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Arun\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Pragadesh\b', re.IGNORECASE), ''),
    (re.compile(r'\bDr\.?\s*Gambira\b', re.IGNORECASE), ''),
    (re.compile(r'\bGopinath\b', re.IGNORECASE), ''),
    (re.compile(r'\bPrakhyath\b', re.IGNORECASE), ''),
    (re.compile(r'\bGambira\b', re.IGNORECASE), ''),
    (re.compile(r'\bPragadesh\b', re.IGNORECASE), ''),
    (re.compile(r'\bMithun\b', re.IGNORECASE), ''),
    (re.compile(r'\bEMP\d+\b', re.IGNORECASE), ''),
]

def sanitize_text(text):
    if not text:
        return text
    clean = text
    for pat, repl in SENSITIVE_PATTERNS:
        clean = pat.sub(repl, clean)
    clean = re.sub(r'[ \t]+', ' ', clean).strip()
    return clean

def sanitize_filename(name):
    clean = sanitize_text(name)
    clean = re.sub(r'[\\/*?:"<>|]', '_', clean)
    clean = re.sub(r'\s+', ' ', clean).strip()
    return clean[:90]

def create_cloned_doc_from_paragraphs(master_doc, start_idx, end_idx, target_path):
    # Create new docx with exact page setup
    new_doc = docx.Document()
    for s_master, s_new in zip(master_doc.sections, new_doc.sections):
        s_new.top_margin = s_master.top_margin
        s_new.bottom_margin = s_master.bottom_margin
        s_new.left_margin = s_master.left_margin
        s_new.right_margin = s_master.right_margin
        s_new.page_width = s_master.page_width
        s_new.page_height = s_master.page_height

    if len(new_doc.paragraphs) > 0:
        p_elem = new_doc.paragraphs[0]._p
        p_elem.getparent().remove(p_elem)

    # Append exact deep-copied <w:p> elements from master
    body = new_doc._body._body
    for i in range(start_idx, end_idx):
        orig_p = master_doc.paragraphs[i]
        p_copy = copy.deepcopy(orig_p._p)
        
        # Sanitize text inside <w:t> tags while keeping all formatting tags (<w:b>, <w:i>, <w:u>, etc.)
        for t_elem in p_copy.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t_elem.text:
                for pat, repl in SENSITIVE_PATTERNS:
                    t_elem.text = pat.sub(repl, t_elem.text)
                    
        body.append(p_copy)

    # Ensure Times New Roman font & 12pt size across all runs
    for p in new_doc.paragraphs:
        for r in p.runs:
            r.font.name = 'Times New Roman'
            if r.font.size is None:
                r.font.size = Pt(12)

    new_doc.save(target_path)

def sanitize_single_docx_file(filepath):
    doc = docx.Document(filepath)
    changed = False
    for p in doc.paragraphs:
        for r in p.runs:
            # Enforce Times New Roman 12
            r.font.name = 'Times New Roman'
            if r.font.size is None:
                r.font.size = Pt(12)
            orig = r.text
            sanitized = orig
            for pat, repl in SENSITIVE_PATTERNS:
                sanitized = pat.sub(repl, sanitized)
            if orig != sanitized:
                r.text = sanitized
                changed = True

    doc.save(filepath)

def get_docx_base64_and_lines(filepath):
    doc = docx.Document(filepath)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    with open(filepath, 'rb') as f:
        b64 = base64.b64encode(f.read()).decode('utf-8')
    return b64, lines
