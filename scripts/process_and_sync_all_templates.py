import os
import re
import json
import base64
import shutil
import docx
from docx.shared import Pt, Inches, RGBColor

DESKTOP_DIR = r'C:\Users\doctor\Desktop\Radiology_Templates_DOCX'
RADNITO_DIR = r'C:\Users\doctor\antigravity\adventurous-babbage\radnito'
RADIOLOGY_DICTATION_DIR = r'C:\Users\doctor\antigravity\adventurous-babbage\radiology-dictation'

def sanitize_filename(name):
    clean = re.sub(r'[\\/*?:"<>|]', '_', name)
    clean = re.sub(r'\s+', ' ', clean).strip()
    return clean[:90]

def create_formatted_docx(filepath, title, paragraphs, font_name="Times New Roman", font_size_pt=12):
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.268)  # A4
        section.page_height = Inches(11.693)

    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    for i, p_text in enumerate(paragraphs):
        text_str = p_text.strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        if not text_str:
            continue

        is_title = (i == 0)
        is_heading = any(text_str.upper().startswith(h) for h in [
            'IMPRESSION:', 'TECHNIQUE:', 'CLINICAL PROFILE:', 'CLINICAL HISTORY:',
            'CLINICAL DIAGNOSIS:', 'OBSERVATIONS:', 'FINDINGS:', 'PAST HISTORY',
            'PROCEDURE:', 'ADVICE:', 'RECOMMENDATION:'
        ])

        if is_title or is_heading:
            r = p.add_run(text_str)
            r.bold = True
            r.font.name = font_name
            r.font.size = Pt(font_size_pt)
        else:
            colon_pos = text_str.find(':')
            if 0 < colon_pos < 25 and not text_str.startswith('http'):
                label = text_str[:colon_pos+1]
                val = text_str[colon_pos+1:]
                r1 = p.add_run(label)
                r1.bold = True
                r1.font.name = font_name
                r1.font.size = Pt(font_size_pt)
                r2 = p.add_run(val)
                r2.font.name = font_name
                r2.font.size = Pt(font_size_pt)
            else:
                r = p.add_run(text_str)
                r.font.name = font_name
                r.font.size = Pt(font_size_pt)

    doc.save(filepath)

def get_docx_base64_and_lines(filepath):
    doc = docx.Document(filepath)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    with open(filepath, 'rb') as f:
        b64 = base64.b64encode(f.read()).decode('utf-8')
    return b64, lines

def main():
    print("=================================================================")
    print("STRICT PROCESSING: CT (RIS-i), MRI A (Dr Prakhyath), MRI B (MRI B.docx)")
    print("=================================================================")

    # 1. Reset Desktop Workspace
    if os.path.exists(DESKTOP_DIR):
        shutil.rmtree(DESKTOP_DIR)
    os.makedirs(DESKTOP_DIR, exist_ok=True)

    ct_dir = os.path.join(DESKTOP_DIR, '01_CT')
    mri_a_dir = os.path.join(DESKTOP_DIR, '02_MRI_A')
    mri_b_dir = os.path.join(DESKTOP_DIR, '03_MRI_B')

    os.makedirs(ct_dir, exist_ok=True)
    os.makedirs(mri_a_dir, exist_ok=True)
    os.makedirs(mri_b_dir, exist_ok=True)

    all_catalog_templates = []
    global_id = 1

    # -----------------------------------------------------------------
    # A. CT FORMATS (Centricity RIS-i Exact Format)
    # -----------------------------------------------------------------
    print("\n--- Processing CT Formats (Centricity RIS-i) ---")
    ct_seen = set()

    # 1. Individual CT Formats from IHSR CT (exact original binary copies)
    ihsr_ct_path = r'[NETWORK_RADIOLOGY_SHARE] imaging Format templates\IHSR Format\CT IHSR'
    if os.path.exists(ihsr_ct_path):
        for f in sorted(os.listdir(ihsr_ct_path)):
            if f.endswith('.docx') and not f.startswith('~$'):
                src = os.path.join(ihsr_ct_path, f)
                dst = os.path.join(ct_dir, f)
                shutil.copy2(src, dst)
                b64, lines = get_docx_base64_and_lines(dst)
                title = lines[0] if lines else os.path.splitext(f)[0]
                norm_title = re.sub(r'[^a-zA-Z0-9]', '', title.lower())
                ct_seen.add(norm_title)
                all_catalog_templates.append({
                    'id': f"ct_{global_id}",
                    'name': title,
                    'title': title,
                    'modality': 'CT',
                    'category': 'CT',
                    'lines': lines,
                    'docxBase64': b64,
                    'content': '\n'.join(lines)
                })
                global_id += 1
                print(f"  [CT] {f} -> {title[:60]}")

    # 2. Master Centricity RIS-i CT collection from CT reporting formats.docx
    prag_ct_path = r'[NETWORK_RADIOLOGY_SHARE] Pragadesh\FORMATS\Radiology report formats\CT reporting formats.docx'
    if os.path.exists(prag_ct_path):
        prag_doc = docx.Document(prag_ct_path)
        current_report_lines = []
        current_report_title = ""
        
        for p in prag_doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue
            is_bold = any(r.bold for r in p.runs)
            is_new_report_header = is_bold and (
                txt.upper().startswith('C.T.') or 
                txt.upper().startswith('CT SCAN') or 
                txt.upper().startswith('HRCT SCAN') or 
                txt.upper().startswith('CT ANGIOGRAM') or 
                txt.upper().startswith('CT ANGIOGRAPHY') or 
                txt.upper().startswith('CT WHOLE BODY')
            ) and len(txt) < 80

            if is_new_report_header:
                if current_report_lines and len(current_report_lines) >= 3:
                    norm = re.sub(r'[^a-zA-Z0-9]', '', current_report_title.lower())
                    if norm not in ct_seen:
                        ct_seen.add(norm)
                        filename = sanitize_filename(current_report_title) + '.docx'
                        out_path = os.path.join(ct_dir, filename)
                        create_formatted_docx(out_path, current_report_title, current_report_lines, font_name="Times New Roman", font_size_pt=12)
                        b64, lines = get_docx_base64_and_lines(out_path)
                        all_catalog_templates.append({
                            'id': f"ct_{global_id}",
                            'name': current_report_title,
                            'title': current_report_title,
                            'modality': 'CT',
                            'category': 'CT',
                            'lines': lines,
                            'docxBase64': b64,
                            'content': '\n'.join(lines)
                        })
                        global_id += 1
                        print(f"  [CT RIS-i] {current_report_title[:60]}")
                
                current_report_title = txt
                current_report_lines = [txt]
            else:
                if current_report_lines:
                    current_report_lines.append(txt)

        if current_report_lines and len(current_report_lines) >= 3:
            norm = re.sub(r'[^a-zA-Z0-9]', '', current_report_title.lower())
            if norm not in ct_seen:
                ct_seen.add(norm)
                filename = sanitize_filename(current_report_title) + '.docx'
                out_path = os.path.join(ct_dir, filename)
                create_formatted_docx(out_path, current_report_title, current_report_lines, font_name="Times New Roman", font_size_pt=12)
                b64, lines = get_docx_base64_and_lines(out_path)
                all_catalog_templates.append({
                    'id': f"ct_{global_id}",
                    'name': current_report_title,
                    'title': current_report_title,
                    'modality': 'CT',
                    'category': 'CT',
                    'lines': lines,
                    'docxBase64': b64,
                    'content': '\n'.join(lines)
                })
                global_id += 1
                print(f"  [CT RIS-i] {current_report_title[:60]}")

    # -----------------------------------------------------------------
    # B. MRI A ([NETWORK_RADIOLOGY_SHARE] Prakhyath Gambira\mri formats)
    # Exclude: Archive formats, *.lnk, *.pptx, fistula.png
    # -----------------------------------------------------------------
    print("\n--- Processing MRI A Formats ([NETWORK_RADIOLOGY_SHARE] Prakhyath Gambira\mri formats) ---")
    mri_a_root = r'[NETWORK_RADIOLOGY_SHARE] Prakhyath Gambira\mri formats'
    ignored_items = {'gopinath formats', 'dr prakhyath gambira - shortcut.lnk', 'sathish-march ppt.pptx', 'gopinath formats - shortcut.lnk', 'fistula.png'}
    mri_a_seen = set()

    for item in sorted(os.listdir(mri_a_root)):
        if item.lower() in ignored_items or item.startswith('~$'):
            continue
        full_path = os.path.join(mri_a_root, item)
        if os.path.isdir(full_path):
            continue  # Ignore subfolders like Archive formats
        
        if item.endswith('.docx'):
            try:
                # Copy exact original binary file directly
                dst_path = os.path.join(mri_a_dir, item)
                shutil.copy2(full_path, dst_path)
                b64, lines = get_docx_base64_and_lines(dst_path)
                if lines:
                    title = lines[0] if len(lines[0]) < 80 and (lines[0].isupper() or lines[0].startswith('MRI') or lines[0].startswith('BRACHIAL') or lines[0].startswith('CARDIAC')) else os.path.splitext(item)[0]
                    norm = re.sub(r'[^a-zA-Z0-9]', '', title.lower())
                    if norm not in mri_a_seen:
                        mri_a_seen.add(norm)
                        all_catalog_templates.append({
                            'id': f"mri_a_{global_id}",
                            'name': title,
                            'title': title,
                            'modality': 'MRI A',
                            'category': 'MRI A',
                            'lines': lines,
                            'docxBase64': b64,
                            'content': '\n'.join(lines)
                        })
                        global_id += 1
                        print(f"  [MRI A] {item} -> {title[:60]}")
            except Exception as e:
                print(f"  [MRI A Warning] Could not process {item}: {e}")

    # -----------------------------------------------------------------
    # C. MRI B (D:\Desktop\MRI B.docx)
    # -----------------------------------------------------------------
    print("\n--- Processing MRI B Formats (D:\\Desktop\\MRI B.docx) ---")
    mri_b_src = r'D:\Desktop\MRI B.docx'
    if os.path.exists(mri_b_src):
        b_doc = docx.Document(mri_b_src)
        
        reports = []
        current_title = ""
        current_lines = []

        for p in b_doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue
            
            is_bold = any(r.bold for r in p.runs)
            is_new_mri_b = is_bold and (
                txt.upper().startswith('MRI SCAN') or 
                txt.upper().startswith('MRI BONY') or 
                txt.upper().startswith('BRACHIAL PLEXUS') or 
                txt.upper().startswith('MRI ') or 
                txt.upper().startswith('MRI -') or 
                txt.upper().startswith('MRI PITUITARY')
            ) and len(txt) < 90

            if is_new_mri_b:
                if current_lines and len(current_lines) >= 3:
                    reports.append((current_title, current_lines))
                current_title = txt
                current_lines = [current_title]
            else:
                if current_lines:
                    current_lines.append(txt)

        if current_lines and len(current_lines) >= 3:
            reports.append((current_title, current_lines))

        print(f"Extracted {len(reports)} distinct reports from MRI B.docx:")
        for idx, (title, lines) in enumerate(reports):
            clean_title = re.sub(r'[:]+$', '', title).strip()
            clean_title = re.sub(r'Clinical Profile.*', '', clean_title, flags=re.IGNORECASE).strip()
            filename = f"{idx+1:02d}_{sanitize_filename(clean_title)}.docx"
            out_path = os.path.join(mri_b_dir, filename)

            # Preserve exact Times New Roman 12pt single spacing
            create_formatted_docx(out_path, clean_title, lines, font_name="Times New Roman", font_size_pt=12)
            b64, file_lines = get_docx_base64_and_lines(out_path)

            all_catalog_templates.append({
                'id': f"mri_b_{global_id}",
                'name': clean_title,
                'title': clean_title,
                'modality': 'MRI B',
                'category': 'MRI B',
                'lines': file_lines,
                'docxBase64': b64,
                'content': '\n'.join(file_lines)
            })
            global_id += 1
            print(f"  [MRI B #{idx+1:02d}] {clean_title}")

    # -----------------------------------------------------------------
    # D. SUMMARY & SAVE TO PROJECTS
    # -----------------------------------------------------------------
    print(f"\nTotal Templates Assembled: {len(all_catalog_templates)}")
    ct_count = len([t for t in all_catalog_templates if t['category'] == 'CT'])
    mri_a_count = len([t for t in all_catalog_templates if t['category'] == 'MRI A'])
    mri_b_count = len([t for t in all_catalog_templates if t['category'] == 'MRI B'])

    print(f"  - CT Templates:    {ct_count}")
    print(f"  - MRI A Templates: {mri_a_count}")
    print(f"  - MRI B Templates: {mri_b_count}")

    # Write README to Desktop
    readme_path = os.path.join(DESKTOP_DIR, 'README_AND_INSTRUCTIONS.txt')
    with open(readme_path, 'w', encoding='utf-8') as f:
        f.write(f"""================================================================================
OFFICIAL RADIOLOGY DOCX TEMPLATES
================================================================================
Workspace Location: C:\\Users\\doctor\\Desktop\\Radiology_Templates_DOCX

Total Active Templates: {len(all_catalog_templates)}
  📁 01_CT     ({ct_count} templates: GE Centricity RIS-i standard CT formats)
  📁 02_MRI_A  ({mri_a_count} templates: MRI_Source MRI formats)
  📁 03_MRI_B  ({mri_b_count} templates: MRI B master formats)

HOW TO USE:
--------------------------------------------------------------------------------
1. EDIT ANY TEMPLATE:
   - Double-click any .docx file to open and edit in Microsoft Word.
   - Exact Times New Roman font, single spacing, and margins are preserved.
   - Save the file (Ctrl+S).

2. ADD NEW TEMPLATES:
   - Paste any new .docx into 01_CT, 02_MRI_A, or 03_MRI_B.

3. SYNC TO RADNITO & RADIOLOGY DICTATION APPS (1-CLICK):
   - Double-click "Sync_Templates_To_App.bat" in this folder.
   - It will automatically package and push everything to GitHub and GitHub Pages!
================================================================================
""")

    # Write sync script and batch file to Desktop
    shutil.copy2(os.path.join(RADNITO_DIR, 'scripts', 'sync_templates_from_desktop.py'), os.path.join(DESKTOP_DIR, 'sync_templates.py'))
    with open(os.path.join(DESKTOP_DIR, 'Sync_Templates_To_App.bat'), 'w') as f:
        f.write('@echo off\ntitle Sync Templates to Apps\npython "%~dp0sync_templates.py"\npause\n')

    # Update radnito & radiology-dictation templatesData.json
    for project_dir in [RADNITO_DIR, RADIOLOGY_DICTATION_DIR]:
        services_dir = os.path.join(project_dir, 'services')
        public_dir = os.path.join(project_dir, 'public')
        os.makedirs(services_dir, exist_ok=True)
        os.makedirs(public_dir, exist_ok=True)

        for p_target in [os.path.join(services_dir, 'templatesData.json'), os.path.join(public_dir, 'templatesData.json')]:
            with open(p_target, 'w', encoding='utf-8') as f:
                json.dump(all_catalog_templates, f, indent=2)
            print(f"[OK] Saved {p_target}")

if __name__ == '__main__':
    main()
