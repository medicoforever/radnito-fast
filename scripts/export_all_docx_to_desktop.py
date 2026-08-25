import os
import re
import json
import base64
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DESKTOP_DIR = r'C:\Users\doctor\Desktop\Radiology_Templates_DOCX'
RADNITO_DIR = r'C:\Users\doctor\antigravity\adventurous-babbage\radnito'
RADIOLOGY_DICTATION_DIR = r'C:\Users\doctor\antigravity\adventurous-babbage\radiology-dictation'

def sanitize_filename(name):
    # Remove invalid Windows filename characters: \ / : * ? " < > |
    clean = re.sub(r'[\\/*?:"<>|]', '_', name)
    clean = re.sub(r'\s+', ' ', clean).strip()
    return clean[:100]

def determine_folder(modality, category, title):
    mod = (modality or '').upper()
    cat = (category or '').upper()
    tit = (title or '').upper()
    comb = f"{mod} {cat} {tit}"

    if 'MRI' in comb or 'MAGNETIC RESONANCE' in comb or 'MRCP' in comb:
        return '01_MRI'
    elif 'CT' in comb or 'COMPUTED TOMOGRAPHY' in comb or 'HRCT' in comb or 'C.T.' in comb:
        return '02_CT'
    elif 'DOPPLER' in comb or 'USG' in comb or 'ULTRASOUND' in comb or 'SONOGRAPHY' in comb:
        return '03_USG_and_Doppler'
    elif 'X-RAY' in comb or 'XRAY' in comb or 'RADIOGRAPH' in comb or 'PA VIEW' in comb:
        return '04_X-Ray'
    elif 'MAMMOGRAM' in comb or 'MAMMOGRAPHY' in comb:
        return '05_Mammography'
    elif any(k in comb for k in ['BARIUM', 'HSG', 'IVU', 'FLUOROSCOPY', 'PROCEDURE', 'BIOPSY', 'ASPIRATION']):
        return '06_Procedures_and_Fluoroscopy'
    else:
        return '07_General_and_Special'

def create_docx_file(filepath, title, content_lines, modality, category):
    doc = docx.Document()
    
    # Page setup - 0.8 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Style: Calibri
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Dark slate

    # Title paragraph
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run(title.upper())
    title_run.bold = True
    title_run.font.size = Pt(15)
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep navy blue

    # Modality & Category subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run(f"Modality: {modality} | Category: {category}")
    sub_run.italic = True
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Slate gray

    # Horizontal divider line (or bottom border)
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_before = Pt(0)
    div_p.paragraph_format.space_after = Pt(12)
    div_run = div_p.add_run("―" * 45)
    div_run.font.size = Pt(8)
    div_run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

    # Body lines
    for line in content_lines:
        line_str = line.strip()
        if not line_str:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)

        # Check for section headers (e.g. FINDINGS:, IMPRESSION:, TECHNIQUE:, CLINICAL INDICATION:)
        header_keywords = [
            'TECHNIQUE:', 'PROTOCOL:', 'FINDINGS:', 'OBSERVATIONS:', 'IMPRESSION:',
            'CLINICAL HISTORY:', 'INDICATION:', 'COMPARISON:', 'CONCLUSION:',
            'PROCEDURE:', 'ADVICE:', 'RECOMMENDATION:'
        ]
        
        is_header = False
        for hk in header_keywords:
            if line_str.upper().startswith(hk):
                is_header = True
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                
                # Split header and rest
                prefix = line_str[:len(hk)]
                suffix = line_str[len(hk):].strip()
                
                r1 = p.add_run(prefix + " ")
                r1.bold = True
                r1.font.size = Pt(11.5)
                r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark
                
                if suffix:
                    r2 = p.add_run(suffix)
                    r2.font.size = Pt(11)
                break

        if not is_header:
            # Bullet point or normal finding paragraph
            if line_str.startswith('-') or line_str.startswith('•') or line_str.startswith('*'):
                p.paragraph_format.left_indent = Inches(0.25)
                clean_bullet_text = re.sub(r'^[-•*]\s*', '', line_str)
                r = p.add_run("•  " + clean_bullet_text)
            else:
                # Check for label like "Liver:", "Gallbladder:", "Spleen:"
                colon_idx = line_str.find(':')
                if 0 < colon_idx < 30 and not line_str.startswith('http'):
                    label = line_str[:colon_idx+1]
                    rest = line_str[colon_idx+1:].strip()
                    r1 = p.add_run(label + " ")
                    r1.bold = True
                    r1.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                    r2 = p.add_run(rest)
                else:
                    r = p.add_run(line_str)

    doc.save(filepath)

def main():
    print(f"Exporting all DOCX templates to {DESKTOP_DIR}...")
    os.makedirs(DESKTOP_DIR, exist_ok=True)

    # 1. Load Primary Catalog Templates (templatesData.json)
    td_path = os.path.join(RADNITO_DIR, 'services', 'templatesData.json')
    primary_templates = []
    if os.path.exists(td_path):
        with open(td_path, 'r', encoding='utf-8-sig') as f:
            primary_templates = json.load(f)
    print(f"Loaded {len(primary_templates)} primary catalog templates from templatesData.json")

    # 2. Load Knowledgebase Templates (report_knowledgebase.json)
    kb_path = os.path.join(RADNITO_DIR, 'public', 'report_knowledgebase.json')
    kb_reports = []
    if os.path.exists(kb_path):
        with open(kb_path, 'r', encoding='utf-8-sig') as f:
            kb_data = json.load(f)
            reports_by_cat = kb_data.get('reports_by_category', {})
            for cat, r_list in reports_by_cat.items():
                for r in r_list:
                    kb_reports.append(r)
    print(f"Loaded {len(kb_reports)} full reports from report_knowledgebase.json")

    # Merge & Deduplicate by normalized title
    seen_titles = set()
    all_templates = []

    # Priority 1: Primary Catalog Templates
    for t in primary_templates:
        name = t.get('name') or t.get('title') or 'Untitled Template'
        norm = re.sub(r'[^a-zA-Z0-9]', '', name.lower())
        if norm not in seen_titles:
            seen_titles.add(norm)
            lines = t.get('lines') or []
            if not lines and t.get('content'):
                lines = t.get('content').split('\n')
            all_templates.append({
                'id': t.get('id', f"tpl_{len(all_templates)+1}"),
                'title': name,
                'modality': t.get('modality') or 'Radiology',
                'category': t.get('category') or 'General',
                'lines': lines,
                'docxBase64': t.get('docxBase64')
            })

    # Priority 2: Knowledgebase Templates
    for r in kb_reports:
        name = r.get('title') or 'Untitled Report'
        norm = re.sub(r'[^a-zA-Z0-9]', '', name.lower())
        if norm not in seen_titles:
            seen_titles.add(norm)
            content = r.get('content') or ''
            lines = [l for l in content.split('\n') if l.strip()]
            all_templates.append({
                'id': r.get('id', f"kb_{len(all_templates)+1}"),
                'title': name,
                'modality': r.get('modality') or 'Radiology',
                'category': r.get('category') or 'General',
                'lines': lines,
                'docxBase64': None
            })

    print(f"Total unique templates to export: {len(all_templates)}")

    stats = {}
    saved_count = 0

    for t in all_templates:
        title = t['title']
        modality = t['modality']
        category = t['category']
        lines = t['lines']
        docx_b64 = t['docxBase64']

        folder_name = determine_folder(modality, category, title)
        target_folder = os.path.join(DESKTOP_DIR, folder_name)
        os.makedirs(target_folder, exist_ok=True)

        filename = sanitize_filename(title) + '.docx'
        filepath = os.path.join(target_folder, filename)

        # Handle filename collisions in same folder
        counter = 1
        base_name = sanitize_filename(title)
        while os.path.exists(filepath):
            filepath = os.path.join(target_folder, f"{base_name}_{counter}.docx")
            counter += 1

        try:
            if docx_b64 and len(docx_b64) > 100:
                # If valid DOCX binary already exists in base64, write it directly
                raw_bytes = base64.b64decode(docx_b64)
                with open(filepath, 'wb') as f:
                    f.write(raw_bytes)
            else:
                # Generate clean formatted docx from lines
                create_docx_file(filepath, title, lines, modality, category)

            saved_count += 1
            stats[folder_name] = stats.get(folder_name, 0) + 1
        except Exception as e:
            print(f"Error saving {filepath}: {e}")

    print(f"\nSuccessfully generated {saved_count} .docx template files!")
    print("Breakdown by folder:")
    for k, v in sorted(stats.items()):
        print(f"  📁 {k}: {v} templates")

    # Create README / Instructions file on Desktop
    readme_path = os.path.join(DESKTOP_DIR, 'README_AND_INSTRUCTIONS.txt')
    with open(readme_path, 'w', encoding='utf-8') as f:
        f.write("""================================================================================
RADIOLOGY DOCX TEMPLATES - DESKTOP REPOSITORY
================================================================================

HOW TO USE THIS FOLDER:
--------------------------------------------------------------------------------
1. EDIT ANY EXISTING TEMPLATE:
   - Open any .docx file in Microsoft Word or LibreOffice.
   - Edit the headings, normal findings, measurements, technique, or impression.
   - Save the file (Ctrl+S).

2. ADD A BRAND NEW TEMPLATE:
   - Simply drop any new .docx file into the corresponding folder
     (e.g., 01_MRI, 02_CT, 03_USG_and_Doppler, 04_X-Ray, etc.) or in 07_General_and_Special.
   - You can name the file whatever you want (e.g. "HRCT Chest Post-COVID.docx").

3. SYNC ALL CHANGES TO RADNITO & RADIOLOGY DICTATION APPS:
   - Double-click "Sync_Templates_To_App.bat" in this folder.
   - It will automatically:
     a) Read all your modified and new .docx files
     b) Package them into the web apps
     c) Build both web apps
     d) Push the updates live to GitHub and GitHub Pages!

================================================================================
Folders:
  📁 01_MRI                        - Brain, Spine, Musculoskeletal, Abdomen, MRCP, Angio
  📁 02_CT                         - Brain, HRCT Chest, Abdomen, Pelvis, Angiography, PNS
  📁 03_USG_and_Doppler            - Abdomen, Pelvis, KUB, Fetal, Thyroid, Carotid, Doppler
  📁 04_X-Ray                      - Chest, Spine, Extremities, Skull, Pelvis, Abdomen
  📁 05_Mammography                - Bilateral Mammogram, Screening, Diagnostic
  📁 06_Procedures_and_Fluoroscopy - Barium, HSG, IVU, Fluoroscopic studies
  📁 07_General_and_Special        - Miscellaneous, Clinical and Special Templates
================================================================================
""")

    print(f"Created instructions at {readme_path}")

if __name__ == '__main__':
    main()
